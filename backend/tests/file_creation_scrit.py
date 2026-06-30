# script for creating test documents to work with

import os
import csv
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    print("pip install python-docx")
    exit(1)

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.lib.units import mm
except ImportError:
    print("pip install reportlab")
    exit(1)

OUTPUT_DIR = Path("tests/fixtures")
PDF_DIR = OUTPUT_DIR / "pdf"
DOCX_DIR = OUTPUT_DIR / "docx"
INVALID_DIR = OUTPUT_DIR / "invalid"

for dir_path in [OUTPUT_DIR, PDF_DIR, DOCX_DIR, INVALID_DIR]:
    try:
        dir_path.mkdir(parents=True, exist_ok=True)
    except Exception as e:
        print(f"Error creating {dir_path}: {e}")
        print("Trying relative path...")
        OUTPUT_DIR = Path("./tests/fixtures")
        PDF_DIR = OUTPUT_DIR / "pdf"
        DOCX_DIR = OUTPUT_DIR / "docx"
        INVALID_DIR = OUTPUT_DIR / "invalid"
        for d in [OUTPUT_DIR, PDF_DIR, DOCX_DIR, INVALID_DIR]:
            d.mkdir(parents=True, exist_ok=True)
        break

def create_valid_pdf(filename, text_content, num_pages=3):
    filepath = PDF_DIR / filename
    c = canvas.Canvas(str(filepath), pagesize=A4)
    width, height = A4
    
    for page in range(num_pages):
        c.setFont("Helvetica-Bold", 16)
        c.drawString(20*mm, height - 20*mm, f"Page {page + 1}: {filename.replace('.pdf', '')}")
        
        c.setFont("Helvetica", 12)
        y_position = height - 40*mm
        words = text_content.split()
        line = ""
        
        for word in words:
            if len(line + " " + word) < 80:
                line += " " + word if line else word
            else:
                c.drawString(20*mm, y_position, line)
                y_position -= 5*mm
                line = word
                if y_position < 20*mm:
                    c.showPage()
                    c.setFont("Helvetica", 12)
                    y_position = height - 40*mm
        
        if line:
            c.drawString(20*mm, y_position, line)
        
        c.setFont("Helvetica", 10)
        c.drawString(width/2 - 10*mm, 15*mm, f"Page {page + 1}")
        
        if page < num_pages - 1:
            c.showPage()
    
    c.save()
    print(f"OK: {filepath}")

pdf_content_1 = """
Le franais est une langue romance parle par environ 300 millions de personnes 
dans le monde. C'est la langue officielle dans 29 pays et l'une des six langues 
officielles des Nations Unies. Le franais est issu du latin vulgaire import dans 
la Gaule par les Romains. La langue a volu au fil des sicles pour devenir le 
franais moderne que nous connaissons aujourd'hui. La grammaire franaise est 
caractrise par des conjugaisons complexes et des rgles d'accord du participe 
pass. Le vocabulaire franais s'est enrichi au cours des sicles, empruntant 
des mots l'italien, l'espagnol, l'anglais et d'autres langues. La francophonie 
est un ensemble de pays et de communauts qui partagent l'usage du franais. 
Les pays francophones sont disperss sur les cinq continents. Le franais est 
galement une langue de culture, de diplomatie et d'affaires. La littrature 
franaise est riche et varie, avec des auteurs comme Victor Hugo, Marcel Proust 
et Albert Camus. La cuisine franaise est rpute dans le monde entier. Le fromage 
et le vin sont des lments importants de la culture franaise. Les rgions de 
France ont chacune leurs spcialits culinaires. La France est le pays le plus 
visit au monde, avec Paris comme destination phare. La Tour Eiffel est le 
monument le plus clbre de France. Le systme ducatif franais est centralis 
et divis en plusieurs cycles. Les coles primaires, les collges et les lyces 
forment le systme scolaire. L'enseignement suprieur comprend les universits 
et les grandes coles. La recherche scientifique est active en France. Le 
CNRS est le plus grand organisme de recherche public en Europe. La France 
est une rpublique avec un prsident lu au suffrage universel. Le Parlement 
est compos de l'Assemble nationale et du Snt. Le systme juridique franais 
est bas sur le droit civil. La Constitution de la Cinquime Rpublique date 
de 1958. La France est membre fondateur de l'Union europenne. L'euro est la 
monnaie officielle depuis 2002. Les symboles de la France incluent le coq, 
le drapeau bleu-blanc-rouge et la Marseillaise. La devise nationale est 
Libert, Egalit, Fraternit. La France est un pays de traditions et de 
modernit. Les Franais sont fiers de leur patrimoine culturel. La langue 
franais est en constante volution avec les nouvelles technologies. Les 
rseaux sociaux influencent de plus en plus le langage. Les jeunes gnrations 
innovent dans la manire de parler franais. Le franais reste une langue 
vivante et dynamique. L'Acadmie franaise est l'institution qui rgule la 
langue franaise. Elle publie le dictionnaire officiel. L'apprentissage du 
franais est populaire dans le monde entier. De nombreuses coles et universits 
enseignent le franais comme langue trangre. Les tests de franais comme le 
DELF et le DALF sont reconnus internationalement. La France accueille de 
nombreux tudiants trangers chaque anne. Paris est une ville universitaire 
de premier plan. La Sorbonne est l'une des universits les plus anciennes 
d'Europe. Le franais est galement parl au Canada, en Belgique, en Suisse 
et en Afrique. Le Qubec est une province francophone au Canada. L'Afrique 
compte le plus grand nombre de francophones au monde. Des pays comme la 
Cte d'Ivoire, le Sngal et le Cameroun ont le franais comme langue officielle. 
La francophonie est en pleine expansion. Les festivals de la francophonie 
c lbrent la diversit culturelle. La chanson franaise est clbre avec des 
artistes comme dith Piaf et Charles Aznavour. Le cinma franais est connu 
pour sa qualit artistique. Les festivals de Cannes sont clbres dans le 
monde entier. La mode franaise est synonyme d'légance et de raffinement. 
Les marques franaises sont rputes dans le monde. Le luxe franais est un 
secteur conomique important. La France est une puissance conomique. Le 
tourisme est un secteur cl de l'conomie franaise. Les Alpes et la 
Mditerrane attirent de nombreux visiteurs. La cte d'Azur est une 
destination touristique pris. Les chteaux de la Loire sont des sites 
historiques remarquables. Le Mont-Saint-Michel est un site class au 
patrimoine mondial. L'architecture franaise est varie et riche. Les 
cathdrales gotiques sont des chefs-d'uvre. Notre-Dame de Paris est 
l'une des plus clbres. Le Louvre est le plus grand muse du monde. 
Il abrite la Joconde de L'onard de Vinci. Le muse d'Orsay est ddi 
l'art du XIXe sicle. L'art contemporain est expos au Centre Pompidou. 
La France est un pays de crativit et d'innovation. Le systme de sant 
franais est considr comme l'un des meilleurs du monde. La scurit 
sociale assure la protection des citoyens. Les Franais sont attachs 
leur systme de protection sociale. Le droit du travail est strict en 
France. Les 35 heures sont la dure lgale du travail. Les syndicats 
sont influents dans la vie sociale. Le dialogue social est important 
dans la culture franaise. Les manifestations sont frquentes en France. 
C'est une tradition de protestation. La Grve est un moyen d'expression 
reconnu. Le pays a une histoire riche et complexe. La Rvolution 
franaise de 1789 a chang le monde. Les droits de l'homme sont un 
hritage important. La Dclaration des droits de l'homme est un texte 
fondateur. Les valeurs rpublicaines sont au cur de la socit franaise. 
La lacit est un principe important. L'glise et l'tat sont spars 
depuis 1905. La libert de conscience est garantie. Les religions 
sont libres en France. La diversit religieuse est une ralit. 
Le judasme, le christianisme et l'islam sont prsents. La cohabitation 
entre les communauts est une question importante. L'intgration est 
un sujet de dbat. La France est une terre d'accueil pour les immigrants. 
Les vagues migratoires ont faonn le pays. La socit franaise est 
multiculturelle. La mixit est une valeur porte par la rpublique. 
L'avenir de la France se construit dans l'Europe et le monde.
"""

pdf_content_2 = """
La grammaire franaise est un systme complexe et structur. Les noms sont 
diviss en genres masculin et fminin. Le pluriel est gnralement form 
en ajoutant un "s" au nom. Les adjectifs s'accordent en genre et en nombre 
avec le nom. Les verbes sont conjugus selon six temps principaux. Le 
prsent, le pass compos, l'imparfait, le plus-que-parfait, le futur simple 
et le conditionnel. Les pronoms personnels sont je, tu, il, elle, nous, 
vous, ils, elles. Les pronoms relatifs sont qui, que, dont, lequel. 
La ngation se fait avec "ne...pas" autour du verbe. Les prpositions 
indiquent les relations spatiales et temporelles. Les conjonctions lient 
les phrases et les ides. L'interrogation peut se faire de plusieurs 
manires. L'inversion du sujet est formelle. Est-ce que est une formule 
courante. Les adverbes modifient le sens des verbes. Les articles 
dfinis sont le, la, les. Les articles indfinis sont un, une, des. 
L'article partitif du, de la, des indique une quantit indtermine. 
Les expressions de quantit utilisent de. La voix passive est forme 
avec l'auxiliaire tre et le participe pass. Le subjonctif est un 
mode utilis pour exprimer le doute ou le souhait. Le conditionnel 
exprime une hypothse ou une condition. L'impratif est utilis pour 
donner des ordres. Le participe prsent est un adjectif verbal. Le 
grondif exprime la simultanit. Les phrases complexes utilisent des 
subordonnes. La proposition relative complte le nom. La proposition 
circonstancielle exprime le temps, la cause ou la consquence. La ponctuation 
est importante pour le sens. La virgule, le point, le point-virgule ont 
leurs rgles. Les guillemets encadrent les citations. Le trait d'union 
est utilis dans certains mots composs. L'orthographe franaise est 
historique et parfois difficile. Les accents aigu, grave et circonflexe 
changent la prononciation. La cdrille donne un son "s" au "c". Le trma 
indique que la voyelle se prononce sparement. Les rgles d'accord du 
participe pass sont complexes. L'accord se fait avec l'auxiliaire avoir 
si le COD est plac avant. Avec tre, le participe s'accorde avec le sujet. 
Les verbes pronominaux ont des rgles spcifiques. Les verbes impersonnels 
comme "il pleut" sont invariables. La concordance des temps est importante 
dans le discours rapport. Le style indirect change les temps et les 
pronoms. Le discours direct utilise les guillemets. La litote est une 
figure de style qui attnue. La mtaphore compare deux ides. La comparaison 
utilise des outils comme "comme". La personnification donne des caractres 
humains. L'hyperbole exagre pour marquer. L'ironie est un procéd de 
style. L'allgorie reprsente une ide abstraite. La littrature franaise 
est pleine de ces figures de style. Les potes utilisent les rimes et 
les mètres. La versification est un art. La prose est plus libre. Les 
diffrents genres littraires incluent le roman, le thtre et la posie. 
Le roman est un genre trs populaire. Le thtre a des tragdies et des 
comdies. La posie explore les mots et les sons. Les mouvements 
littraires comme le classicisme, le romantisme et le surralisme ont 
marqu l'histoire. Le classicisme recherche l'harmonie et la clart. 
Le romantisme exprime l'motion et l'imagination. Le surralisme explore 
l'inconscient et les rves. Les auteurs franais sont tudis dans le 
monde entier. La littrature franaise est un pilier de la culture mondiale.
"""

pdf_content_3 = """
La civilisation franaise est riche en traditions et en innovations. La 
gastronomie franaise est reconnue comme un patrimoine culturel immatriel. 
Le repas gastronomique des Franais est inscrit au patrimoine de l'UNESCO. 
Les fromages franais sont varis et nombreux. On compte plus de 400 sortes 
de fromages en France. Le pain est un lment essentiel du repas. La 
baguette est un symbole de la France. Les vins franais sont les meilleurs 
du monde. Les rgions comme Bordeaux, Bourgogne et Champagne sont clbres. 
La champagne est une rgion productrice de vin ptillant. Le champagne est 
servi lors des grandes occasions. Les huîtres sont un plat de ftes. 
Les escargots sont une spcialit franaise. Le foie gras est un produit 
de luxe. La cuisine franaise est base sur des produits frais. Les 
marchs sont des lieux de vie importants. Le pain, le fromage et le vin 
forment un repas simple mais complet. Les habitudes alimentaires voluent 
avec le temps. La restauration rapide est devenue courante. Mais la cuisine 
traditionnelle reste apprcie. Les chefs franais sont clbres dans le monde. 
La cuisine molculaire est une innovation franaise. Les coles de cuisine 
franaises forment les meilleurs chefs. La mode franaise est un secteur 
d'excellence. Paris est la capitale de la mode. Les couturiers franais 
comme Chanel et Dior sont des lgends. Le luxe franais est une industrie 
majeure. Les boutiques de luxe sont sur les Champs-lyses. Le prt-porter 
est aussi important. Les marques de mode franaises sont internationales. 
Le design franais est reconnu pour son esthtique. Les meubles franais 
sont des oeuvres d'art. Le style Louis XV est un classique. L'architecture 
franais est varige. Les chteaux de la Loire sont des joyaux. Le palais 
de Versailles est un symbole de la monarchie. Le Louvre est un ancien 
palais royal. La tour Eiffel est une construction spectaculaire. Le 
centre Pompidou est un exemple d'architecture moderne. La France est 
un pays d'art et de culture. Les muses franais sont parmi les plus 
visits. Le muse du Louvre est le plus visit du monde. Le muse d'Orsay 
est ddi l'art impressionniste. Les expositions temporaires attirent 
les foules. Les festivals sont nombreux en France. Le festival de Cannes 
est le plus clbre. Le festival d'Avignon est ddi au thtre. Les 
festivals de musique comme les Eurockennes sont populaires. Les Franais 
aiment les ftes et les clbrations. Le 14 juillet est la fte nationale. 
Il y a des feux d'artifice dans toute la France. La fte de la musique a 
lieu le 21 juin. Tous les musiciens jouent dans la rue. Les marchs de 
Nol sont des traditions. L'Alsace a des marchs de Nol clbres. Les 
lampes de Nol sont une tradition dans le sud. La galette des rois est 
mange en janvier. Le mois de janvier est celui des galettes. Le chocolat 
chaud est une boisson d'hiver. Les bistros sont des lieux de convivialit. 
Les cafs sont des rendez-vous importants. Les terrasses sont pleines l't. 
Le monde des affaires est dynamique en France. Les grandes entreprises 
franaises sont prsentes dans le monde. Le CAC 40 est l'indice boursier. 
L'conomie est diversifie entre industrie et services. L'agriculture est 
aussi importante. La France est le premier producteur agricole d'Europe. 
Le bl, le vin et le fromage sont les principaux produits. L'industrie 
automobile a des marques comme Renault et Peugeot. L'arme française 
est moderne et professionnelle. La France est une puissance nuclaire. 
Le porte-avions Charles de Gaulle est un symbole de puissance. La marine 
et l'arme de l'air sont aussi prsentes. La diplomatie franaise est 
active dans le monde. La France est membre permanent du Conseil de 
scurit de l'ONU. La francophonie est un instrument d'influence. Les 
relations internationales sont un enjeu majeur. La France entretient des 
liens avec ses anciennes colonies. La coopration est importante avec 
l'Afrique. L'aide au dveloppement est une priorité. La culture franaise 
rayonne travers le monde. La langue franaise est un vecteur de valeurs. 
L'ducation est une priorit nationale. L'cole publique est laïque et 
gratuite. Les coles prives sont aussi prsentes. Les tudiants franais 
font des tudes suprieures. Les universits accueillent de nombreux 
tudiants. Les grandes coles forment l'lite. La recherche est un moteur 
d'innovation. Les scientifiques franais sont reconnus. Le prix Nobel a 
t attribu plusieurs Franais. Les mdicaments franais sont utiliss dans 
le monde. La sant est un secteur de qualit. Les hpitaux franais sont 
performants. Le systme de retraite est solidaire. Les Franais vivent 
longtemps et en bonne sant. L'esprance de vie est leve. La qualit de 
vie est bonne en France. Les transports sont dvelopps. Le TGV est le 
train grande vitesse franais. La France a un rseau autoroutier dense. 
Les aroports internationaux desservent le monde. Le service public est 
un pilier de la rpublique. La poste, les tlcommunications et l'nergie 
sont des services essentiels. La France investit dans les nergies 
renouvelables. L'nergie nuclaire fournit l'lectricit. La transition 
cologique est en cours. Les Franais sont sensibles l'environnement. 
Le tri des dchets est pratique courante. Les voitures lectriques se 
dveloppent. La France est un pays de progress et de tradition. L'avenir 
est prometteur pour la France dans l'Europe unie.
"""

create_valid_pdf("french_civilization.pdf", pdf_content_1, num_pages=5)
create_valid_pdf("french_grammar.pdf", pdf_content_2, num_pages=4)
create_valid_pdf("french_culture.pdf", pdf_content_3, num_pages=5)

def create_valid_docx(filename, text_content):
    filepath = DOCX_DIR / filename
    doc = Document()
    
    title = filename.replace('.docx', '').replace('_', ' ').title()
    doc.add_heading(title, level=1)
    
    paragraphs = text_content.split('\n\n')
    for para in paragraphs:
        if para.strip():
            doc.add_paragraph(para.strip())
    
    doc.add_heading('Vocabulary', level=2)
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    table.cell(0, 0).text = 'French'
    table.cell(0, 1).text = 'English'
    table.cell(1, 0).text = 'Bonjour'
    table.cell(1, 1).text = 'Hello'
    table.cell(2, 0).text = 'Merci'
    table.cell(2, 1).text = 'Thank you'
    table.cell(3, 0).text = 'Oui'
    table.cell(3, 1).text = 'Yes'
    
    doc.save(str(filepath))
    print(f"OK: {filepath}")

docx_content_1 = """
Le franais est une langue vivante et dynamique. L'apprentissage du franais 
est accessible tous. Il existe de nombreuses mthodes pour apprendre le 
franais. Les cours de franais sont dispenses dans le monde entier. Les 
ressources en ligne sont nombreuses et varies. Les applications mobiles 
facilitent l'apprentissage. Le franais est une langue cl pour la culture 
et le voyage. La France est une destination touristique majeure. Le 
franais est parl dans de nombreux pays. La variété des accents enrichit 
la langue. Le franais de France diffre du franais du Qubec. Le franais 
d'Afrique a ses propres expressions. La comprhension est possible entre 
tous les francophones. La communication est l'objectif principal. Parler 
franais ouvre des portes professionnelles. Les entreprises recherchent 
des candidats bilingues. La connaissance du franais est un atout. Les 
relations internationales utilisent le franais. Les organisations 
internationales emploient le franais. L'ONU a le franais comme langue 
de travail. Les traités et accords sont souvent rdigs en franais. La 
diplomatie franaise est active. Les sommets francophones sont r guliers. 
La communaut francophone est solidaire. La diversité est une force. Les 
cultures francophones sont riches. Les arts et la littrature franais 
sont apprcis. La chanson franaise a un public fidle. Le cinma 
franaise est cratif. Les spctacles francophones attirent des publics 
varis. Les festivals francophones sont des vnements importants. La 
francophonie est aussi une organisation politique. L'OIF regroupe les 
tats francophones. La coopration est un objectif majeur. Les projets 
communs sont nombreux. La jeunesse francophone est engage. Les coles 
et universits participent. Les changes tudiants sont frquents. 
L'apprentissage du franais est une aventure humaine. Chaque locuteur 
contribue la vitalit de la langue. Le franais est une langue d'avenir. 
Elle s'adapte aux changes du monde. Les nouvelles technologies intgrent 
le franais. Les rseaux sociaux ont des communauts francophones. La 
langue volue constamment. Les jeunes crent de nouveaux mots. Le 
verlan est un exemple d'volution. Les anglicismes sont aussi prsents. 
Le franais reste une langue de rfrence. Son apprentissage est une 
richesse. Parler franais est un voyage dans le temps. C'est aussi une 
ouverture sur le monde. La Francophonie est une communaut vivante. 
Ensemble, les francophones construisent l'avenir. La langue franaise 
est un trsor partager. Chacun peut apporter sa pierre l'difice. 
L'importance du franais ne se dment pas. Il est une langue de culture 
et de modernit. La France est un pays cl de la francophonie. Mais 
la francophonie dpasse les frontires de la France. Elle est prsente 
sur tous les continents. Chaque region apporte sa couleur. Cette 
diversit est une chance. Le franais est donc une langue mlange. 
Un mlange de traditions et d'innovations. Un mlange de cultures. 
Un mlange de peuples. Voil la richesse du franais.
"""

docx_content_2 = """
Les bases de la grammaire franaise sont essentielles. La phrase simple 
a un sujet et un verbe. Le sujet peut tre un nom ou un pronom. Le verbe 
indique l'action ou l'tat. Le complment complte le sens de la phrase. 
Les types de phrases sont : dclarative, interrogative, exclamative et 
imprative. La phrase dclarative donne une information. La phrase 
interrogative pose une question. La phrase exclamative exprime une 
motion. La phrase imprative donne un ordre ou un conseil. La structure 
de la phrase est gnralement SVO. Sujet - Verbe - Objet. Les adverbes 
peuvent modifier le verbe. Ils indiquent le temps, le lieu ou la manire. 
Les prpositions introduisent des complments. Elles sont importantes 
pour la comprhension. Les conjonctions lient les mots et les ides. 
Il y a des conjonctions de coordination et de subordination. Les 
conjonctions de coordination sont : mais, ou, et, donc, or, ni, car. 
Les conjonctions de subordination introduisent des propositions. Les 
pronoms remplacent les noms pour viter les rp titions. Les pronoms 
personnels sujets sont indispensables. Les pronoms complments sont 
plus complexes. Ils peuvent tre directs ou indirects. Les pronoms 
relatifs relient les propositions. Qui, que, dont, lequel sont les 
principaux. Le choix du pronom relatif dpend de la fonction. L'accord 
du participe pass est une rgle difficile. Le participe pass avec 
l'auxiliaire avoir s'accorde avec le COD si plac avant. Avec tre, 
il s'accorde avec le sujet. Les verbes pronominaux ont des rgles 
spcifiques. Le participe pass s'accorde avec le sujet sauf si le 
verbe est suivi d'un COD. Les rgles sont nombreuses mais logiques. 
L'orthographe franaise a ses particularits. Les accents sont importants. 
L'accent aigu sur le "e" donne le son "". L'accent grave sur le "e" 
donne le son "". L'accent circonflexe peut tre sur toutes les voyelles. 
Le trma indique la prononciation spare. La cdrille transforme le "c" 
en "s" devant "a", "o", "u". Les rgles de l'orthographe sont fixes. 
Il faut les apprendre par cœur. La maitrise de la grammaire est un 
avantage. Elle permet de s'exprimer correctement. La pratique est 
essentielle. Lire en franais amliore la grammaire. crire en franais 
consolide les acquis. Parler avec des francophones est la meilleure 
mthode. Les erreurs font partie de l'apprentissage. Il ne faut pas 
les craindre. La persvrance est la qualit principale. La grammaire 
franais est accessible avec de la pratique. Chaque jour, il faut 
s'exercer. Les progrès sont visibles rapidement. La confiance 
vient avec la pratique. Parler franais est un plaisir. La grammaire 
est la structure de la langue. La comprendre est fondamental. Elle 
permet de bien communiquer. La communication est le but ultime. La 
grammaire en est l'outil. Utilisons-la pour mieux parler et crire. 
La langue franaise est belle et prcise. Sa grammaire reflte cette 
prcision. Apprenons-la avec enthousiasme. La matrise de la grammaire 
franais est la cl de la russite. Elle ouvre les portes de la culture 
et de la connaissance. Le franais est une langue de rfrence. 
Savoir le parler correctement est un atout. La grammaire en est 
la base solide. Construisons sur cette base. L'avenir est prometteur. 
Le franais continue d'voluer. Sa grammaire s'adapte. Restons 
attentifs aux volutions. L'apprentissage est un processus continu. 
Chaque jour apporte son lot de dcouvertes. La grammaire franaise 
est un sujet fascinant. Explorons-la ensemble.
"""

docx_content_3 = """
Les rgles de conjugaison franaise sont nombreuses. Les verbes sont 
diviss en trois groupes. Le premier groupe comprend les verbes en 
-er. Le deuxime groupe comprend les verbes en -ir avec un participe 
prsent en -issant. Le troisime groupe comprend tous les autres verbes. 
Les verbes du premier groupe sont rguliers. Ils suivent un modle simple. 
Le prsent des verbes en -er est facile. Je parle, tu parles, il parle. 
Nous parlons, vous parlez, ils parlent. Les verbes du deuxime groupe 
sont aussi rguliers. Je finis, tu finis, il finit. Nous finissons, 
vous finissez, ils finissent. Le troisime groupe est irrgulier. Il 
faut apprendre les conjugaisons par coeur. tre et avoir sont les 
auxiliaires. Ils sont indispensables pour les temps composs. Le pass 
compos se forme avec l'auxiliaire et le participe pass. Je suis all 
pour tre. J'ai mang pour avoir. Le choix de l'auxiliaire est important. 
Les verbes de mouvement utilisent tre. Les verbes transitifs utilisent 
avoir. Les verbes pronominaux utilisent tre. Le futur simple exprime 
un vnement futur. Je parlerai, tu parleras, il parlera. Le conditionnel 
exprime une hypothse. Je parlerais si j'avais le temps. L'imparfait 
dcrit des actions passes continues. Je parlais quand tu es arriv. 
Le plus-que-parfait exprime une action passe avant une autre. J'avais 
parl avant de partir. Le subjonctif exprime le doute et le souhait. 
Il faut que je parle. Le mode impratif exprime un ordre. Parle! 
Les participes prsents et passs sont importants. Le participe prsent 
est un adjectif. Le participe pass est utilis dans les temps composs. 
Les verbes irrguliers les plus courants sont : aller, venir, faire, 
dire, voir, pouvoir, vouloir, devoir, savoir. Ils sont trs utiliss. 
Il est essentiel de les connatre. Les verbes permettent d'exprimer 
toutes les actions. Sans verbes, la communication est impossible. 
La conjugaison est la base de la grammaire. Elle est apprendre dès 
le dbut. La pratique rgulire est ncessaire. Les exercices de 
conjugaison sont utiles. Les jeux peuvent faciliter l'apprentissage. 
La mémorisation des formes est importante. Les rgles sont des outils. 
Il faut les utiliser sans cesse. L'oral et l'crit sont complmentaires. 
La conjugaison est partout dans la langue. Chaque phrase contient un 
verbe. Le verbe est le cœur de la phrase. Il en est le centre. Bien 
conjuguer est fondamental. Les Franais sont exigeants sur ce point. 
Une mauvaise conjugaison est mal vue. Mais les erreurs sont pardonnables. 
L'important est de progresser. La motivation est la cl. La pratique 
avec des natifs est excellente. Les changes linguistiques sont 
bénéfiques. Les cours de franais sont aussi importants. Ils apportent 
des bases solides. Les manuels sont des outils de rfrence. Les 
exercices sont varis et adapts. La difficult progresse avec le temps. 
Il ne faut pas se dcourager. Chaque personne apprend son rythme. 
La persvrance paie toujours. Un jour, la conjugaison deviendra naturelle. 
Elle ne sera plus un obstacle. Elle sera un outil au service de 
l'expression. Alors, on pourra communiquer librement. C'est l'objectif 
de tout apprentissage. La libert d'expression est la plus belle rcompense.
"""

create_valid_docx("french_language_learning.docx", docx_content_1)
create_valid_docx("french_grammar_basics.docx", docx_content_2)
create_valid_docx("french_conjugation.docx", docx_content_3)

def create_empty_pdf():
    filepath = INVALID_DIR / "empty.pdf"
    with open(filepath, 'wb') as f:
        pass
    print(f"OK: {filepath}")

def create_broken_pdf():
    filepath = INVALID_DIR / "broken.pdf"
    with open(filepath, 'wb') as f:
        f.write(b'%PDF-1.4\n')
        f.write(b'broken content\x00\x01\x02\x03\xff\xfe\xfd\xfc\n')
        f.write(b'%%EOF\n')
        f.write(b'\x00' * 1000)
    print(f"OK: {filepath}")

def create_empty_docx():
    filepath = INVALID_DIR / "empty.docx"
    doc = Document()
    doc.save(str(filepath))
    print(f"OK: {filepath}")

def create_wrong_format_files():
    filepath = INVALID_DIR / "image_as_pdf.pdf"
    with open(filepath, 'wb') as f:
        f.write(b'\xFF\xD8\xFF\xE0' + b'\x00' * 100)
    print(f"OK: {filepath}")
    
    filepath = INVALID_DIR / "txt_as_docx.docx"
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write("This is a text file, not DOCX\n" * 10)
    print(f"OK: {filepath}")
    
    filepath = INVALID_DIR / "too_large.txt"
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write("x" * 25 * 1024 * 1024)
    print(f"OK: {filepath}")

def create_docx_with_bad_font():
    filepath = INVALID_DIR / "bad_font.docx"
    doc = Document()
    doc.add_heading('Bad Font Test', level=1)
    
    p = doc.add_paragraph('This text uses Comic Sans MS Bold font')
    run = p.runs[0]
    run.font.name = 'Comic Sans MS Bold'
    run.font.size = Pt(14)
    
    p2 = doc.add_paragraph('This text uses Impact font')
    run2 = p2.runs[0]
    run2.font.name = 'Impact'
    run2.font.size = Pt(16)
    
    p3 = doc.add_paragraph('Special symbols:   ')
    run3 = p3.runs[0]
    run3.font.name = 'Arial Unicode MS'
    
    doc.save(str(filepath))
    print(f"OK: {filepath}")

def create_edge_case_files():
    filepath = PDF_DIR / "many_pages.pdf"
    c = canvas.Canvas(str(filepath), pagesize=A4)
    for i in range(20):
        c.drawString(20*mm, 250*mm, f"Page {i+1} of 20")
        c.drawString(20*mm, 230*mm, "French text for large document processing")
        c.showPage()
    c.save()
    print(f"OK: {filepath}")
    
    filepath = PDF_DIR / "unicode_pdf.pdf"
    c = canvas.Canvas(str(filepath), pagesize=A4)
    c.setFont("Helvetica", 14)
    y = 250*mm
    texts = [
        "Bonjour tout le monde!",
        "Test:   ",
        "French accents: é è ê ë",
        "Symbols:     "
    ]
    for text in texts:
        c.drawString(20*mm, y, text)
        y -= 10*mm
    c.save()
    print(f"OK: {filepath}")

def create_precision_dataset():
    filepath = OUTPUT_DIR / "precision_queries.csv"
    queries = [
        ["query", "expected_file", "expected_chunk_text"],
        ["grammaire", "french_grammar.pdf", "grammaire franaise"],
        ["conjugaison", "french_conjugation.docx", "rgles de conjugaison"],
        ["civilisation", "french_civilization.pdf", "civilisation franaise"],
        ["culture", "french_culture.pdf", "culture franaise"],
        ["verbe", "french_grammar.pdf", "verbes sont conjugus"],
        ["subjonctif", "french_conjugation.docx", "Le subjonctif"],
        ["gastronomie", "french_civilization.pdf", "gastronomie franaise"],
        ["littrature", "french_culture.pdf", "littrature franaise"],
        ["apprentissage", "french_language_learning.docx", "apprentissage du franais"],
        ["francophonie", "french_culture.pdf", "francophonie"],
    ]
    
    with open(filepath, 'w', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        writer.writerows(queries)
    print(f"OK: {filepath}")

create_empty_pdf()
create_broken_pdf()
create_empty_docx()
create_wrong_format_files()
create_docx_with_bad_font()
create_edge_case_files()
create_precision_dataset()

print("\n" + "="*60)
print("ALL FILES CREATED SUCCESSFULLY")
print(f"FOLDER: {OUTPUT_DIR.absolute()}")
print("="*60)

pdf_files = list(PDF_DIR.glob("*.pdf"))
docx_files = list(DOCX_DIR.glob("*.docx"))
invalid_files = list(INVALID_DIR.glob("*"))

print(f"\nStatistics:")
print(f"   Valid PDF: {len(pdf_files)}")
print(f"   Valid DOCX: {len(docx_files)}")
print(f"   Invalid files: {len(invalid_files)}")
print(f"   Total: {len(pdf_files) + len(docx_files) + len(invalid_files)} files")